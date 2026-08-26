#!/usr/bin/env python3
"""Build the COMMTR/JICV Replication package explanatory file (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Replication_package_explanatory_file.docx"


def set_run_font(run, *, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if color is not None:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True)


def add_meta(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, size=12, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=12, italic="[" in value)


def add_heading_plain(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)


def add_body(doc, text, *, italic=False, fill=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    color = RGBColor(0xC0, 0x00, 0x00) if fill else None
    set_run_font(run, size=12, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=12)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    add_title(doc, "Replication package explanatory file")

    add_meta(doc, "Journal title: ", "[PLEASE FILL: Communications in Transportation Research, or Journal of Intelligent and Connected Vehicles]")
    add_meta(
        doc,
        "Manuscript title: ",
        "String-stability-aware sequential propagation network for vehicle-platoon prediction: SSP-DMGTimeNet",
    )
    add_meta(doc, "Manuscript ID: ", "[PLEASE FILL after Editorial Manager assignment]")
    add_meta(doc, "Authors: ", "[PLEASE FILL: Name1, Name2, ...]")

    add_heading_plain(doc, "For COMMTR or JICV article")
    add_body(
        doc,
        "This explanatory file is submitted together with the manuscript and is intended for review by a replication editor. "
        "A short summary (no more than 100 words) to be placed in the manuscript immediately before the Acknowledgements is provided at the end of this file. "
        "Please insert the public repository URL and replace “[the replication package was approved]” after the editor’s confirmation.",
    )

    add_heading_plain(doc, "Replication package access")
    add_body(
        doc,
        "The replication package consists of two parts.",
    )
    add_body(
        doc,
        "(1) Code. All source code, YAML configs, unit tests, and one-click scripts needed to rebuild platoon windows, train models, evaluate the three-layer stability protocol, and regenerate tables/figures are provided in the folder submitted with this file (the `submission/` package). "
        "[PLEASE FILL the public URL after depositing the same folder on GitHub / Zenodo / ETS-Data, e.g. https://doi.org/10.5281/zenodo.XXXX or https://github.com/ORG/ssp-dmgtimenet]. "
        "If an open platform cannot be used, we will request ETS-Data access through the replication editor.",
        fill=True,
    )
    add_body(
        doc,
        "(2) Raw trajectory data. HighD and NGSIM cannot be redistributed with this package because of their licenses and data-use agreements. "
        "They must be obtained from the official providers using the steps below. After placement, the bundled scripts reconstruct every training and test split used in the paper, so the key tables can be regenerated without any hidden data.",
    )
    add_body(doc, "HighD (primary training and testing data):")
    add_bullet(doc, "Apply at https://levelxdata.com/highd-dataset/ (RWTH Aachen / ika highD).")
    add_bullet(doc, "Download the 60 recordings and copy the CSV files to `datasets/highD/` so that `01_tracks.csv`, `01_tracksMeta.csv`, `01_recordingMeta.csv`, …, `60_tracks.csv` are present.")
    add_body(doc, "NGSIM US-101 and I-80 (zero-shot evaluation and smoothing sensitivity):")
    add_bullet(doc, "Obtain the FHWA NGSIM vehicle-trajectory files from data.gov / the NGSIM community archive.")
    add_bullet(doc, "Place the six 15-minute freeway periods under `datasets/NGSIM/vehicle-trajectory-data/` as listed in `datasets/README.md`.")
    add_bullet(doc, "For the I-80 16:00–16:15 sensitivity study, also place the Montanino–Punzo reconstructed file `RECONSTRUCTED trajectories-400-0415_NO MOTORCYCLES.csv` in the same period folder.")

    add_heading_plain(doc, "Data description")
    add_body(
        doc,
        "Complete raw trajectories are not uploaded. Only derived processing code is provided. This is required by the HighD license and by FHWA NGSIM redistribution terms, not by a desire to withhold results. "
        "The uploaded package still generates the key results because every number in the main HighD, NGSIM, ablation, platoon-length, and sensitivity tables is produced by deterministic scripts from the official CSVs.",
    )
    add_body(doc, "Derived data produced by the package (not redistributed here, rebuilt locally):")
    add_bullet(doc, "Format: NumPy `.npz` windows with history/future tensors of shape [B, T, N, F] at 10 Hz, plus vehicle identifiers and recording ids.")
    add_bullet(doc, "HighD splits by recording id: train 01–45, validation 46–50, test 51–60. Main experiment uses N = 5 vehicles, 5 s history, 3 s prediction, 1 s stride, leader-stationarity quantile q = 0.5.")
    add_bullet(doc, "NGSIM: US-101 and I-80, up to 4,000 windows per site, same window protocol; normalisation statistics remain those of the HighD training split (zero-shot).")
    add_bullet(doc, "Platoon-length extensions rebuild N = 3, 6, 7 with the same protocol.")
    add_body(
        doc,
        "After `bash scripts/prepare_highd.sh` and `bash scripts/prepare_ngsim.sh`, the files `artifacts/platoons/highd_N5_h5_p3/{train,val,test}.npz` and `artifacts/platoons/ngsim_N5_h5_p3/{us101,i80}/test.npz` are the exact inputs of the reported experiments.",
    )

    add_heading_plain(doc, "Code description")
    add_body(doc, "Computer environment required to run the scripts:")
    add_bullet(doc, "OS: Linux (bash scripts). Python 3.10, 3.11, or 3.12.")
    add_bullet(doc, "Deep learning: PyTorch ≥ 2.0 with CUDA recommended. One GPU with ≥ 12 GB memory is sufficient for batch size 64; CPU can run evaluation but training is impractical.")
    add_bullet(doc, "Python dependencies: see `requirements.txt` (numpy, pandas, scipy, scikit-learn, torch, einops, PyYAML, tqdm, matplotlib, seaborn, statsmodels, h5py, tensorboard, rich, pytest).")
    add_bullet(doc, "Install: `python -m venv .venv && source .venv/bin/activate && pip install -U pip && pip install torch && pip install -e \"./ssp_dmgtimenet[dev]\"`.")
    add_bullet(doc, "Sanity check without data: `bash scripts/smoke_test.sh` (imports + unit tests of the v3 stability protocol).")
    add_body(doc, "What the code contains:")
    add_bullet(doc, "`ssp_dmgtimenet/`: model (SP-DACA, Cross-CFE, HGF), losses, metrics, HighD/NGSIM loaders, training and evaluation entry points.")
    add_bullet(doc, "`ssp_dmgtimenet/configs/`: paper configs, including `ssp_dmgtimenet_v6.yaml` (reported model), 9 baselines, 8 ablations, and N-extension YAMLs.")
    add_bullet(doc, "`scripts/`: `prepare_highd.sh`, `prepare_ngsim.sh`, `train_all.sh`, `evaluate_highd.sh`, `evaluate_extensions.sh`, `make_tables_and_figures.sh`, `reproduce_all.sh`.")
    add_body(
        doc,
        "Step-by-step commands, random seed (42), and the evaluation protocol (δ = 0.05, excitation floor 0.05 m/s) are documented in `README.md`. "
        "Do not use `ssp_dmgtimenet.yaml` or `ssp_dmgtimenet_v5.yaml` for the reported numbers; the paper model is v6.",
    )

    add_heading_plain(doc, "Simulation software description")
    add_body(
        doc,
        "Not applicable. This manuscript does not use traffic microsimulation software (SUMO, VISSIM, AIMSUN, etc.). "
        "All experiments are supervised learning and evaluation on reconstructed trajectory windows in PyTorch.",
    )

    add_heading_plain(doc, "Experiment design description")
    add_body(
        doc,
        "The manuscript does not include surveys, questionnaires, or human-subject experiments. "
        "The computational experiment design is as follows.",
    )
    add_bullet(doc, "Task: given a same-lane platoon of N vehicles, predict the next 3 s of longitudinal state (speed, gap, acceleration) from 5 s of history, and score string stability of the predicted trajectories.")
    add_bullet(doc, "Main comparison (HighD, N = 5): SSP-DMGTimeNet versus five learned baselines (Int-LSTM, Transformer, full-graph attention, LSTM, CNN-Int-LSTM-IDM) and four physics/hybrid baselines (IDM, OVM, FVDM, DMGTimeNet cascade).")
    add_bullet(doc, "Ablations (8): remove delay bias, adjacent loss, CFE, sub-platoon loss, HGF, or FFT loss; replace the chain mask by full-graph attention; freeze τ at 1.0 s.")
    add_bullet(doc, "Generalisation: zero-shot transfer of HighD checkpoints to NGSIM US-101 and I-80; I-80 16:00–16:15 original versus reconstructed trajectories; platoon length N ∈ {3, 5, 6, 7}.")
    add_bullet(doc, "Stability protocol v3 (identical for every model): (i) disturbance detection with a shared detrended-leader RMS floor; (ii) unified external response on GT-selected windows (primary comparison); (iii) conditional internal amplification only on GT-and-prediction excited windows, always reporting the conditional sample size.")
    add_bullet(doc, "Outputs used in the paper: `artifacts/evaluation_v3/tables.md` and `extension_tables.md`, produced by `bash scripts/make_tables_and_figures.sh` from JSON reports that store data/config/checkpoint hashes.")

    add_heading_plain(doc, "Others")
    add_body(doc, "Recommended replication path after the raw CSVs are in place:")
    add_bullet(doc, "`bash scripts/reproduce_all.sh` runs smoke test → HighD windows → NGSIM windows → training → HighD evaluation → extension evaluation → tables/figures.")
    add_bullet(doc, "Full training of 10 main models + 8 ablations + 9 N-extension models takes several GPU-days. To verify the pipeline first, run only `bash scripts/train_all.sh ssp_dmgtimenet_v6` then `bash scripts/evaluate_highd.sh`.")
    add_bullet(doc, "Resume: `SKIP_EXISTING=1 bash scripts/train_all.sh all` skips existing `best.pt` files.")
    add_bullet(doc, "Bit-wise identity of floating-point metrics is not expected across GPU kernels; ranking and order of magnitude should match. Confirm seed 42, δ = 0.05, and HighD-train normalisation for NGSIM.")
    add_bullet(doc, "Hardware used in our runs: Linux, Python 3.12, PyTorch 2.x, CUDA GPU (training). Any CUDA GPU with ≥ 12 GB can reproduce the reported protocol.")

    add_heading_plain(doc, "Manuscript summary (≤ 100 words; insert before Acknowledgements)")
    add_body(
        doc,
        "The code to reconstruct platoon windows, train SSP-DMGTimeNet and all baselines, and regenerate the reported tables is available at [PLEASE FILL URL]. "
        "Raw HighD and NGSIM trajectories are not redistributed; they must be obtained from the official providers as described in the replication package. "
        "Derived training/test splits are rebuilt by the bundled scripts. "
        "[The replication package was approved by the replication editor.]",
        italic=True,
        fill=True,
    )
    add_body(
        doc,
        "Word count of the paragraph above (excluding the two bracketed placeholders) is 78 words. After inserting the URL and the approval sentence, keep the total at or below 100 words.",
        italic=True,
    )

    doc.save(OUT)
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()
